import math
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from gptils import *

############################################
# 2) 메인: 보고서 자동생성 함수
############################################

def generate_rc_report(excel_path, output_docx='RC_Design_Report.docx'):
    """
    1) Excel(SectionData 시트)에 부재별 입력값(b,h,d,fck,fy,Vu,Mu 등)을 적어둔다.
    2) 본 함수를 호출하면 각 부재별 휨·전단·사용성 검토를 수행,
       Word(docx) 형태 보고서를 자동 작성한다.
    """
    # (1) 엑셀 불러오기
    df = pd.read_excel(excel_path, sheet_name='SectionData')

    # 보고서 생성
    doc = Document()
    heading = doc.add_heading('철근콘크리트 단면검토 자동보고서', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p_intro = doc.add_paragraph(
    #     "본 보고서는 Python 스크립트를 이용해 자동 작성되었으며, "
    #     "직사형 단면 + 단철근 가정하에 휨·전단·사용성 검토를 시연한 예시입니다.\n"
    #     "설계기준(KDS/KCI/ACI)에 따라 식과 계수는 수정 필요할 수 있습니다."
    # )

    for _, row in df.iterrows():
        member_id = row['MemberID']
        b   = row['b(mm)']
        h   = row['h(mm)']
        d   = row['d(mm)']  # 이미 유효깊이로 기입했다고 가정
        fck = row['fck(MPa)']
        fy  = row['fy(MPa)']
        Vu  = row['Vu(kN)']
        Mu  = row['Mu(kN·m)']

        # -----------------------------
        #  휨검토
        # -----------------------------
        As_req = required_rebar_area(Mu, b, d, fck, fy)
        if As_req is None:
            # 단면이 너무 작아 해 없음
            As_use = None
        else:
            # 예시: 필요철근량 대비 15% 여유
            As_use = As_req * 1.15

        # 최소·최대철근
        As_min, As_max = calc_min_max_rebar(b, d, fck, fy,
                                            As_req=As_req,
                                            As_use=As_use)

        # 사용률
        usage_ratio = 0.0
        is_as_ok = False
        if (As_req is not None) and (As_use is not None):
            usage_ratio = As_use / As_req
            if (As_use >= As_min) and (As_use <= As_max):
                is_as_ok = True

        # 중립축 깊이
        c, cmax, is_c_ok = (0,0,False)
        if (As_use is not None):
            c, cmax, is_c_ok = check_neutral_axis_depth(As_use, b, d, fck, fy)

        # 철근 변형률
        eps_s, eps_yd = (0.0, 0.0)
        is_yield = False
        if (As_use is not None):
            eps_s = compute_tensile_strain(d, c, eps_cu=0.0033)
            phi_s = 0.90
            Es = 200000  # MPa
            eps_yd = (phi_s * fy) / Es
            is_yield = (eps_s >= eps_yd)
        
        # 설계 휨강도 산정
        Mn, phi_Mn = calculate_design_flexural_strength(As_use, b, d, c, fck, fy, phi=0.90)
        safety_factor = phi_Mn / Mu
        safety_pass = safety_factor >= 1.0

        # -----------------------------
        #  전단검토
        # -----------------------------
        fvy = 400
        Av_use, s = 506.8, 125
        phi, phi_c = 0.9, 0.65
        
        # shear_ok, Vc, Vs = calc_shear_check(Vu, b, d, fck, fy)
        Vcd, Vcd_min, shear_pass = shear_check(Vu, b, d, fck, phi_c)
        pv_min, pv_use, min_shear_pass = min_shear_rebar(Av_use, s, b, fck, fvy)
        Smax = shear_spacing_check(d)
        delta_Tr, delta_T, additional_pass = additional_tension(Mn, Mu, Vu, d)
        ft, fct, crack_pass = serviceability_check(Mu, b, h, fck)


        # -----------------------------
        #  사용성(균열응력) 검토
        # -----------------------------
        # 비균열 가정시 연단인장응력 ft = M / Z
        # 여기서 Z ≈ b * h^2 / 6 (직사형)
        # Z = b*(h**2)/6.0  # mm^3
        # Mu_Nmm = Mu * 1e6 # N·mm
        # ft = Mu_Nmm / Z   # N/mm² = MPa
        # # fct ~ 0.23*(fck^(2/3)) (간단 추정)
        # fct = 0.23*(fck**(2/3))
        # no_crack = (ft <= fct)

         # 사용성검토
        Z = b * h**2 / 6
        ft = Mu * 1e6 / Z
        fct = 0.23 * fck**(2/3)
        stress_check = ft <= fct

        As_cr_min = 0.002 * b * h
        indirect_crack_check = As_use >= As_cr_min

        w_max = (d - c) * eps_s * 0.6  # 예시 간접 균열폭 계산
        allowable_crack_width = 0.3
        crack_width_check = w_max <= allowable_crack_width

        # -----------------------------
        #  보고서 작성
        # -----------------------------
        doc.add_heading(f"부재 ID: {member_id}", level=2)

        p = doc.add_paragraph()
        p.add_run("■ 단면치수 및 재료특성\n").bold = True
        p.add_run(
            f" - 폭 b = {b} mm, 전체높이 h = {h} mm, 유효깊이 d = {d} mm\n"
            f" - fck = {fck} MPa, fy = {fy} MPa\n"
            f" - Vu = {Vu:.2f} kN, Mu = {Mu:.2f} kN·m\n"
        )

        # ─ 휨검토 ─
        p.add_run("\n[1] 휨검토\n").bold = True
        if As_req is None:
            p.add_run(" - 필요철근량(As_req) 계산 불가 → 단면 부족으로 추정\n")
        else:
            p.add_run(f" - 필요한 철근량 As_req = {As_req:.1f} mm²\n")
            p.add_run(f" - 가정 사용철근량 As_use = {As_use:.1f} mm² (약 {usage_ratio*100:.1f}% 여유)\n")
            p.add_run(f"   * 최소철근량(As_min) = {As_min:.1f} mm²\n")
            p.add_run(f"   * 최대철근량(As_max) = {As_max:.1f} mm²\n")
            p.add_run(f"   → 철근량 검토: {'OK' if is_as_ok else 'NG'}\n")

            # 중립축 및 철근 변형률
            p.add_run(f" - 중립축 깊이 c = {c:.1f} mm, cmax = {cmax:.1f} mm → {'OK' if is_c_ok else 'NG'}\n")
            p.add_run(
                "   (식: c = (φs·As_use·fy) / (α·φc·0.85·fck·b))\n"
                f"   (여기서 cmax = 0.4·d = {0.4*d:.1f} mm 가정)\n"
            )

            p.add_run(f" - 인장철근 변형률 εs = {eps_s:.4f}, 항복변형률 εyd = {eps_yd:.4f}\n")
            p.add_run(f"   → 철근항복여부: {'항복' if is_yield else '비항복'}\n")

            p.add_run(f" - 설계 휨강도 산정 Mr = {phi_Mn:.4f}, O.K [ratio= {safety_factor if safety_pass else None :.4f}]\n")
            

        # ─ 전단검토 ─
        p.add_run("\n[2] 전단검토\n").bold = True
        # p.add_run(
        #     f" - 무근전단강도 Vc = {Vc:.2f} kN, 필요 전단보강 Vs = {Vs:.2f} kN\n"
        #     f"   → 전단검토 결과: {'OK' if shear_ok else 'NG'}\n"
        #     "   (식 예: Vc = 0.53·√fck·b·d·10^-3 [kN])\n"
        # )
        p.add_run(f"전단저항(Vcd): {Vcd / 1e3:.3f}kN, 최소값(Vcd_min): {Vcd_min / 1e3:.3f}kN, 검토: {'O.K' if shear_pass else 'N.G'}")
        p.add_run(f"최소철근량(pv_min): {pv_min:.6f}, 사용철근량(pv_use): {pv_use:.6f}, 검토: {'O.K' if min_shear_pass else 'N.G'}")
        p.add_run(f"최대간격(Smax): {Smax:.2f}mm, 실제간격: {s}mm, 검토: {'O.K' if Smax >= s else 'N.G'}")
        p.add_run(f"추가인장력(ΔTr): {delta_Tr:.3f}kN, 작용추가인장력(ΔT): {delta_T:.3f}kN, 추가철근필요: {'불필요' if additional_pass else '필요'}")


        # ─ 사용성(균열응력) ─
        p.add_run("\n[3] 사용성(균열) 검토\n").bold = True
        # p.add_run(
        #     f" - 단순 비균열 가정 시 인장연단응력 ft = {ft:.2f} MPa\n"
        #     f" - fct(추정) = {fct:.2f} MPa → 균열{'미발생' if no_crack else '발생가능'}\n"
        # )
        # p.add_run(
        #     "   (식 예: ft = (Mu×10^6) / (b·h²/6), fct ≈ 0.23·(fck^(2/3)))\n"
        # )
        doc.add_paragraph(
            f"비균열 가정시 인장 연단 응력(ft): {ft:.2f}MPa, 균열강도(fct): {fct:.2f}MPa - {'O.K' if stress_check else 'N.G'}\n"
            f"균열제어 최소철근량(As_cr_min): {As_cr_min:.2f}mm², 사용철근량: {As_use:.2f}mm² - {'O.K' if indirect_crack_check else 'N.G'}\n"
            f"간접 균열폭 계산: {w_max:.3f}mm, 허용 균열폭: {allowable_crack_width:.3f}mm - {'O.K' if crack_width_check else 'N.G'}"
        )

        
        p.add_run("─" * 60 + "\n")

    # (3) 보고서 저장
    doc.save(output_docx)
    print(f"[완료] 보고서가 생성되었습니다: {output_docx}")


############################################
# 3) 실행 예시
############################################
if __name__ == "__main__":
    input_excel = "section_input_template.xlsx"  # 예: b,h,d,fck,fy,Vu,Mu 등
    output_file = "RC_Design_Report.docx"

    generate_rc_report(input_excel, output_docx=output_file)
    print("모든 계산 완료!")
