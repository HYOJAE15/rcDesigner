import math
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from gptils_2 import *

# -----------------------------------------------------
# 메인 보고서 생성 함수
# -----------------------------------------------------
def generate_rc_report(excel_path, output_docx='RC_Design_Report.docx'):
    df = pd.read_excel(excel_path, sheet_name='SectionData')

    doc=Document()
    doc.styles['Normal'].font.name='돋움'
    doc.styles['Normal'].font.size=Pt(10)

    title=doc.add_heading('철근콘크리트 단면검토 보고서',0)
    title.alignment=WD_ALIGN_PARAGRAPH.CENTER

    for idx,row in df.iterrows():
        # 입력값들
        member_id=row['MemberID']
        b=row['b(mm)']
        h=row['h(mm)']
        d=row['d(mm)']
        cover=row.get('cover(mm)',0)
        fck=row['fck(MPa)']
        fy=row['fy(MPa)']
        Vu=row['Vu(kN)']
        Mu=row['Mu(kN·m)']
        # Av_use=row.get('Av(mm²)',0)
        # s=row.get('s(mm)',0)
        Av_use=506.8
        s=125

        doc.add_heading(f'{member_id}',level=1)

        # ㉮ 단면제원 및 설계조건
        p=doc.add_paragraph()
        p.add_run('㉮ 단면제원 및 설계가정\n').bold=True
        p.add_run(f' - b={b} mm, h={h} mm, d={d} mm, cover={cover} mm\n'
                  f' - fck={fck} MPa, fy={fy} MPa\n'
                  f' - Mu={Mu:.2f} kN·m, Vu={Vu:.3f} kN\n')

        # ㉯ 필요철근량 산정
        As_req=required_rebar_area(Mu,b,d,fck,fy)
        p.add_run('\n㉯ 필요철근량 산정\n').bold=True
        if As_req is None:
            p.add_run(' → 이차방정식 해 없음 => 단면부족 => ∴ N.G\n')
            As_use=0.0
        else:
            p.add_run(f' - 필요철근량 As_req= {As_req:.3f} mm² ')
            # p.add_run('∴ O.K\n')

            # 사용철근량
            As_use=As_req*1.15
            usage_ratio=1.15
            p.add_run(f' - 사용철근량= {As_use:.3f} mm² (사용률: {usage_ratio:.3f})\n')

        # ㉰ 철근량 검토
        p.add_run('\n㉰ 철근량 검토\n').bold=True
        As_min,As_max=calc_min_max_rebar(b,d,fck,fy,As_req,As_use)
        if As_use>=As_min and As_use<=As_max:
            p.add_run(f' - As_min= {As_min:.1f} ≤ 사용철근= {As_use:.1f} ≤ As_max= {As_max:.1f} => ∴ O.K\n')
        else:
            p.add_run(f' - As_min= {As_min:.1f}, As_max= {As_max:.1f}, 사용철근= {As_use:.1f} => ∴ N.G\n')

        # ㉱ 중립축 깊이 검토
        p.add_run('\n㉱ 중립축 깊이 검토\n').bold=True
        c, cmax, c_ok=(0,0,False)
        if As_use>0:
            c, cmax, c_ok=check_neutral_axis_depth(As_use,b,d,fck,fy)
            p.add_run(f' - Cmax= {cmax:.1f} mm, c= {c:.1f} mm => {"∴ O.K" if c_ok else "∴ N.G"}\n')

        # ㉲ 인장철근 변형률
        p.add_run('\n㉲ 인장철근 변형률\n').bold=True
        eps_s, eps_yd=(0.0,0.0)
        is_yield=False
        if As_use>0:
            eps_s=compute_tensile_strain(d,c,0.0033)
            eps_yd=(0.9*fy)/200000
            is_yield=(eps_s>=eps_yd)
            p.add_run(f' - εs= {eps_s:.4f}, εyd= {eps_yd:.4f} => {"항복" if is_yield else "비항복"}\n')

        # ㉳ 설계 휨강도 산정
        p.add_run('\n㉳ 설계 휨강도 산정\n').bold=True
        Mn_kNm,phi_Mn=calculate_design_flexural_strength(As_use,b,d,c,fck,fy)
        ratio=(phi_Mn/Mu) if Mu>0 else 0.0
        p.add_run(f' - φMn= {phi_Mn:.3f} kN·m / Mu= {Mu:.3f} => [Ratio : {ratio:.3f}] => {"∴ O.K" if ratio>=1.0 else "∴ N.G"}\n')

        # ================== 전단검토 ====================
        p=doc.add_paragraph()
        p.add_run('【 전단에 대한 검토 】\n').bold=True

        # Vcd check
        Vcd,Vcd_min,shear_ok=shear_check(Vu,b,d,fck,phi_c=0.65)
        p.add_run(f' - Vcd= {Vcd/1e3:.3f} kN, Vcd_min= {Vcd_min/1e3:.3f}, Vu= {Vu:.3f} => {"∴ O.K" if shear_ok else "∴ N.G"}\n')

        # 최소 전단철근량
        pv_min,pv_use,pv_ok=min_shear_rebar(Av_use,s,b,fck,400) # fvy=400
        p.add_run(f' - 최소 전단철근량 검토: pv_min= {pv_min:.4f}, pv_use= {pv_use:.4f} => {"∴ O.K" if pv_ok else "∴ N.G"}\n')

        # 전단철근 간격
        smax=shear_spacing_check(d)
        p.add_run(f' - 전단철근 간격 Smax= {smax:.1f} mm / 실제= {s} => {"∴ O.K" if smax>=s else "∴ N.G"}\n')

        # 추가 종방향 인장력
        dTr,dT, add_ok=additional_tension(Mn_kNm,Mu,Vu,d)
        p.add_run(f' - 추가 종방향 인장력: ΔTr= {dTr:.3f} kN, ΔT= {dT:.3f} kN => {"추가철근 불필요" if add_ok else "추가철근 필요"}\n')

        # ================== 사용성검토(균열) ====================
        p=doc.add_paragraph()
        p.add_run('【 균열 단면의 사용성 검토 】\n').bold=True

        ft,fct,crack_ok=serviceability_check(Mu,b,h,fck)
        p.add_run(f' - 비균열 가정시 인장 연단응력 ft= {ft:.3f} MPa, fct= {fct:.3f} MPa => {"∴ O.K" if crack_ok else "∴ N.G"}\n')

        # 균열제어 최소철근(예시)
        As_cr_min=0.002*b*h
        crack_rebar_ok=(As_use>=As_cr_min)
        p.add_run(f' - 균열제어 최소철근량= {As_cr_min:.1f} mm² / 사용철근= {As_use:.1f} => {"∴ O.K" if crack_rebar_ok else "∴ N.G"}\n')

        # 간접 균열폭
        w_calc=0.0
        if As_use>0:
            eps_s=compute_tensile_strain(d,c,0.0033)
            w_calc=(d-c)*eps_s*0.6
        w_allow=0.30
        w_ok=(w_calc<=w_allow)
        p.add_run(f' - 간접 균열폭 계산= {w_calc:.3f} mm / 허용= {w_allow:.3f} => {"∴ O.K" if w_ok else "∴ N.G"}\n')

        # 구분선
        doc.add_paragraph("─"*80)

    doc.save(output_docx)
    print(f"[완료] 보고서가 생성되었습니다: {output_docx}")

# 실행 예시
if __name__=="__main__":
    input_excel="section_input_template.xlsx"
    output_docx="RC_Design_Report_2.docx"
    generate_rc_report(input_excel, output_docx)
    print("모든 계산 완료!")
