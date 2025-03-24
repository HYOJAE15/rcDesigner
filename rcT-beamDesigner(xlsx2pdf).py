import pandas as pd
from docx import Document
from docx.shared import Inches

from utils import *

def generate_rc_report(excel_path, output_docx='RC_Report.docx'):
    """
    1. 엑셀 파일(설계 입력)을 읽는다
    2. 각 부재별 휨·전단 검토를 수행한다
    3. docx 형태의 보고서를 생성한다
    """
    # (1) 엑셀 데이터 불러오기
    df = pd.read_excel(excel_path, sheet_name='SectionData')

    # 보고서(docx) 문서 생성
    doc = Document()
    doc.add_heading('철근콘크리트 단면검토 자동보고서', level=1)
    doc.add_paragraph('본 문서는 Python 스크립트를 통해 생성된 자동화 예시 보고서입니다.')
    
    # (2) 각 행(부재)에 대해 반복 계산
    for idx, row in df.iterrows():
        member_id = row['MemberID']
        b = row['b(mm)']
        h = row['h(mm)']
        d = row['d(mm)']
        cover = row['cover(mm)']
        fck = row['fck(MPa)']
        fy  = row['fy(MPa)']
        Vu  = row['Vu(kN)']
        Mu  = row['Mu(kN·m)']
        
        ###################
        ### (a) 휨 검토 ###
        ###################

        ## 필요철근량 산정
        As_required = solve_rebar_area_for_flexure(Mu, b, d, fck, fy)
        # 사용 철근량(가정). 여기서는 예시로 “필요량보다 15% 많은 값”이라고 가정
        # 실제는 사용자가 별도 셀에 입력하거나, 철근 호칭 지름·EA 조합으로 산출해야 함
        if As_required is None:
            As_use = None
            ratio_flex = 999  # 또는 Fail
        else:
            As_use = As_required * 1.15  
            
        ## 철근량 검토
        As_min, As_max = calc_min_max_rebar(b, d, fck, fy, As_req=As_required, As_use=As_use)
        usage_ratio = As_use / As_required
        if As_use >= As_min and As_use <= As_max:
            As_ok = True
        else:
            As_ok = False

        ## 중립축 깊이 검토
        c, cmax, c_ok = check_neutral_axis_depth(As_use, b, d, fck, fy)

        ## 인장철근 변형률
        phi_s = 0.90
        Es = 200000
        eps_s = compute_tensile_strain(d, c, eps_cu=0.0033)
        eps_yd = (phi_s * fy) / Es  # MPa / (MPa/unitstrain) = unitstrain
        is_yield = (eps_s >= eps_yd)
        

        
        #####################
        ### (b) 전단 검토 ###
        #####################

        shear_ok, Vc, Vs = calc_shear_check(Vu, b, d, fck, fy)
        
        
        
        
        # (c) 사용성(균열폭 등) 검토
        
        Ec = 200000  # 단위 MPa, 단순 가정
        fct = 0.23 * (fck**(2/3))  # (예시) 장기값 고려 안 함
        # 비균열 가정시 극단섬유인장응력 ft ~ M/Z 로 가정
        # Z ~ b*h^2/6 (직사단면 기준), 단순예시  ≈≈
        Z = b*(h**2)/6.0
        ft = (Mu*1e6)/Z  # [N/mm^2] = MPa
        crack_check = (ft <= fct)
        
        
        
        
        # 보고서에 기재할 문단 생성
        doc.add_heading(f'부재 ID: {member_id}', level=2)

        para = doc.add_paragraph()
        para.add_run(f"단면치수: 폭 b = {b} mm, 높이 h = {h} mm, 유효깊이 d = {d} mm\n").bold=True
        para.add_run(f"콘크리트강도 fck = {fck} MPa, 철근강도 fy = {fy} MPa\n")
        para.add_run(f"설계 전단력 Vu = {Vu:.3f} kN, 설계 휨모멘트 Mu = {Mu:.3f} kN·m\n")

        # 휨결과
        if As_required is None:
            para.add_run("[휨검토] 요구철근량 계산 불가 → 단면이 너무 작은 것으로 추정\n").italic = True
        else:
            para.add_run(f"[휨검토]\n")
            para.add_run(f"필요철근량 As_req = {As_required:.3f} mm²\n")
            para.add_run(f"사용철근량 As_use = {As_use:.3f} mm² → 사용률 = {usage_ratio:.3f}\n")
            para.add_run(f"철근량검토 결과: {'OK' if As_ok else 'NG'}\n")

            para.add_run(f"중립축 깊이 c = {c:.3f} mm\n")
            para.add_run(f"허용 cmax = {cmax:.3f} mm → 사용률 = {'OK' if c_ok else 'NG'}\n")
            
            para.add_run(f"인장철근 변형률 eps_s = {eps_s:.3f}\n")
            para.add_run(f"항복변형률 eps_yd = {eps_yd:.3f} → 항복가정 {'OK' if is_yield else 'NG'}\n")
            
            

        
        
        # 전단결과
        para.add_run(f"[전단검토] Vc(무근전단강도 추정) = {Vc:.1f} kN, Vs(필요 전단보강) = {Vs:.1f} kN\n")
        para.add_run(f"            전단검토 결과: {'OK' if shear_ok else 'NG'}\n")
        
        # 균열여부/사용성
        if crack_check:
            para.add_run(f"[사용성/균열] 비균열 상태로 추정 (ft={ft:.2f} MPa ≤ fct={fct:.2f} MPa)\n")
        else:
            para.add_run(f"[사용성/균열] 균열 발생 가능 (ft={ft:.2f} MPa > fct={fct:.2f} MPa)\n")
            para.add_run(" → KDS 24 14 21 4.2.3.4 식 등에 따른 균열폭 계산 및 최소철근량 확인 필요\n")
        
        para.add_run("--------------------------------------------------------------------------------\n")

    # (3) 최종 docx 파일 저장
    doc.save(output_docx)
    print(f"보고서 생성 완료: {output_docx}")


# ---------------------------
# 3) 실제 스크립트 실행 (직접 호출 예시)
# ---------------------------

if __name__ == "__main__":
    # 준비된 엑셀 파일 경로
    input_excel = "section_input_template.xlsx"  
    output_file = "RC_Design_Report.docx"

    generate_rc_report(input_excel, output_docx=output_file)
    print("모든 계산 완료!")
